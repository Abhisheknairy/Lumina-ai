from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import RedirectResponse
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
import os
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
import io
import docx
from PyPDF2 import PdfReader

# Required for local testing without HTTPS
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'

app = FastAPI(title="Google Drive RAG Backend")

# Configuration for OAuth
CLIENT_SECRETS_FILE = "client_secret.json" 
SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
REDIRECT_URI = "http://localhost:8000/auth/callback"

# In-memory session storage 
user_sessions = {}
# Store the flow temporarily so we don't lose the security verifier
oauth_flows = {} 

from fastapi.middleware.cors import CORSMiddleware

app = FastAPI(title="Google Drive RAG Backend")

# --- CORS BLOCK ---
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
# ---------------------------

# 1. Initialize the OAuth Flow
@app.get("/login")
def login(user_id: str):
    """Generates the Google Login URL and redirects the user."""
    flow = Flow.from_client_secrets_file(
        CLIENT_SECRETS_FILE, scopes=SCOPES, redirect_uri=REDIRECT_URI
    )
    
    auth_url, state = flow.authorization_url(
        access_type='offline',
        prompt='consent',
        state=user_id 
    )
    
    # Store the EXACT flow object in memory using the state as the key
    oauth_flows[state] = flow
    return RedirectResponse(url=auth_url)

# 2. Handle the Google Callback
@app.get("/auth/callback")
def auth_callback(state: str, code: str):
    """Google redirects here. We use the stored flow to trade the code for tokens."""
    # Retrieve the original flow object
    flow = oauth_flows.get(state)
    if not flow:
        raise HTTPException(status_code=400, detail="Session expired or invalid state")
    
    # Trade the code for credentials (it now remembers the verifier!)
    flow.fetch_token(code=code)
    credentials = flow.credentials
    
    # Save the credentials
    user_sessions[state] = {
        'token': credentials.token,
        'refresh_token': credentials.refresh_token,
        'token_uri': credentials.token_uri,
        'client_id': credentials.client_id,
        'client_secret': credentials.client_secret,
        'scopes': credentials.scopes
    }
    
    # Clean up the temporary flow
    del oauth_flows[state]
    
    # Redirect the user BACK to the React frontend Dashboard
    return RedirectResponse(url=f"http://localhost:5173/chat?user_id={state}")

# 3. Fetch Files from a Specific Folder
@app.get("/fetch-folder/{user_id}/{folder_id}")
def fetch_folder_contents(user_id: str, folder_id: str):
    """Uses the user's stored token to list files in a specific Google Drive folder."""
    if user_id not in user_sessions:
        raise HTTPException(status_code=401, detail="User not authenticated")
        
    from google.oauth2.credentials import Credentials
    creds = Credentials(**user_sessions[user_id])
    
    try:
        service = build('drive', 'v3', credentials=creds)
        
        query = f"'{folder_id}' in parents and trashed=false"
        results = service.files().list(
            q=query, 
            fields="nextPageToken, files(id, name, mimeType)",
            pageSize=100
        ).execute()
        
        files = results.get('files', [])
        
        return {
            "folder_id": folder_id,
            "total_files": len(files),
            "files": files
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# Initialize Embeddings & Vector Store Directory
embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
CHROMA_PERSIST_DIR = "./chroma_db"

# 4. Ingest and Vectorize Folder Contents
@app.post("/ingest-folder/{user_id}/{folder_id}")
def ingest_folder_contents(user_id: str, folder_id: str):
    """Downloads files, extracts text, chunks it, and saves to ChromaDB."""
    if user_id not in user_sessions:
        raise HTTPException(status_code=401, detail="User not authenticated")
        
    from google.oauth2.credentials import Credentials
    creds = Credentials(**user_sessions[user_id])
    service = build('drive', 'v3', credentials=creds)
    
    # 1. Fetch the file list
    query = f"'{folder_id}' in parents and trashed=false"
    results = service.files().list(q=query, fields="files(id, name, mimeType)").execute()
    files = results.get('files', [])
    
    if not files:
        return {"message": "No files found in this folder."}

    all_documents = []
    
    # 2. Download and Extract Text
    for f in files:
        file_id = f['id']
        file_name = f['name']
        mime_type = f['mimeType']
        text_content = ""
        
        try:
            if mime_type == 'application/vnd.google-apps.document':
                # Export Google Docs as plain text
                request = service.files().export_media(fileId=file_id, mimeType='text/plain')
                text_content = request.execute().decode('utf-8')
                
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Download Word Docs
                request = service.files().get_media(fileId=file_id)
                content = request.execute()
                doc = docx.Document(io.BytesIO(content))
                text_content = "\n".join([para.text for para in doc.paragraphs])
                
            elif mime_type == 'application/pdf':
                # Download PDFs
                request = service.files().get_media(fileId=file_id)
                content = request.execute()
                pdf_reader = PdfReader(io.BytesIO(content))
                for page in pdf_reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text_content += extracted + "\n"
                        
            # If we successfully extracted text, store it with metadata
            if text_content.strip():
                all_documents.append({
                    "text": text_content, 
                    "metadata": {"source": file_name, "folder_id": folder_id, "user_id": user_id}
                })
                
        except Exception as e:
            print(f"Failed to process {file_name}: {str(e)}")

    # 3. Chunk the Text
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000, 
        chunk_overlap=150,
        separators=["\n\n", "\n", " ", ""]
    )
    
    chunks = []
    metadatas = []
    
    for doc in all_documents:
        split_texts = text_splitter.split_text(doc["text"])
        chunks.extend(split_texts)
        metadatas.extend([doc["metadata"]] * len(split_texts))
        
    # 4. Save to Persistent ChromaDB
    if chunks:
        vector_store = Chroma.from_texts(
            texts=chunks,
            embedding=embeddings,
            metadatas=metadatas,
            persist_directory=CHROMA_PERSIST_DIR
        )
        return {
            "message": "Ingestion Complete!", 
            "files_processed": len(all_documents),
            "total_chunks_saved": len(chunks)
        }
    else:
        return {"message": "Files were found, but no readable text could be extracted."}

from pydantic import BaseModel
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate

# Ensure your Google API Key is set here
GOOGLE_API_KEY = "AIzaSyAv2I5cUmNXI88c8_79xRexfoMa7kFvcS8"

# Initialize Gemini LLM
llm = ChatGoogleGenerativeAI(
    model="gemini-2.5-flash",
    google_api_key=GOOGLE_API_KEY,
    temperature=0.3
)

# Define the structure of our incoming request
class ChatRequest(BaseModel):
    question: str

# Store chat histories
chat_histories = {} 

# 5. The Core RAG Chat Endpoint
@app.post("/chat/{user_id}/{folder_id}")
def chat_with_documents(user_id: str, folder_id: str, request: ChatRequest):
    """Searches the vector DB, utilizes chat history, and answers the user."""
    
    # 1. Retrieve or initialize the user's chat history
    if user_id not in chat_histories:
        chat_histories[user_id] = []
    
    # Format history into a readable string
    history_text = "\n".join([f"{msg['role']}: {msg['content']}" for msg in chat_histories[user_id]])
    
    # 2. Connect to ChromaDB
    vector_store = Chroma(
        persist_directory=CHROMA_PERSIST_DIR,
        embedding_function=embeddings
    )
    
    # 3. Search for context
    docs = vector_store.similarity_search(
        query=request.question,
        k=4, 
        filter={"folder_id": folder_id} 
    )
    
    context_text = "\n\n---\n\n".join([doc.page_content for doc in docs]) if docs else "No relevant documents found."
    
    # 4. Create the prompt
    prompt_template = ChatPromptTemplate.from_template("""
    You are a highly intelligent and helpful AI assistant. 
    
    You have access to the user's document context and their previous conversation history.
    - If the user asks a question about their documents, prioritize answering from the "Document Context".
    - If the user asks a general question (like "Hi", "Write me a python script", etc.), answer using your general knowledge.
    - Always be conversational and acknowledge the previous chat history if relevant.
    
    Previous Conversation History:
    {history}
    
    Document Context:
    {context}
    
    User Question:
    {question}
    """)
    
    prompt = prompt_template.format(
        history=history_text, 
        context=context_text, 
        question=request.question
    )
    
    # 5. Get the response from Gemini
    response = llm.invoke(prompt)
    
    # 6. Save this interaction to memory
    chat_histories[user_id].append({"role": "Human", "content": request.question})
    chat_histories[user_id].append({"role": "AI", "content": response.content})
    
    # Keep memory from getting too huge (keep last 10 messages)
    if len(chat_histories[user_id]) > 10:
        chat_histories[user_id] = chat_histories[user_id][-10:]
    
    sources = list(set([doc.metadata.get("source") for doc in docs])) if docs else []
    
    return {
        "question": request.question,
        "answer": response.content,
        "sources_used": sources,
        "chat_history_length": len(chat_histories[user_id]) // 2
    }

def get_all_files_recursive(service, item_id):
    """Helper function to find every file within a folder and its sub-folders."""
    files_found = []
    
    # Query for items inside the current folder
    query = f"'{item_id}' in parents and trashed=false"
    page_token = None
    
    while True:
        results = service.files().list(
            q=query,
            corpora="allDrives",
            includeItemsFromAllDrives=True,
            supportsAllDrives=True,
            fields="nextPageToken, files(id, name, mimeType)",
            pageSize=1000,
            pageToken=page_token
        ).execute()
        
        items = results.get('files', [])
        
        for item in items:
            if item['mimeType'] == 'application/vnd.google-apps.folder':
                # Recursive call: If we find a sub-folder, go inside it
                files_found.extend(get_all_files_recursive(service, item['id']))
            else:
                # It's a file, add it to our list
                files_found.append(item)
        
        page_token = results.get('nextPageToken')
        if not page_token:
            break
            
    return files_found

@app.post("/ingest-item/{user_id}/{item_id}")
def ingest_item(user_id: str, item_id: str):
    """Recursively downloads every file in a folder tree and ingests it."""
    if user_id not in user_sessions:
        raise HTTPException(status_code=401, detail="User not authenticated")
        
    from google.oauth2.credentials import Credentials
    creds = Credentials(**user_sessions[user_id])
    service = build('drive', 'v3', credentials=creds)
    
    # 1. Get metadata of the root item selected
    root_item = service.files().get(
        fileId=item_id, 
        supportsAllDrives=True, 
        fields="id, name, mimeType"
    ).execute()
    
    files_to_process = []
    
    # 2. Determine if we need to dive deep
    if root_item['mimeType'] == 'application/vnd.google-apps.folder':
        print(f"Deep searching folder: {root_item['name']}...")
        files_to_process = get_all_files_recursive(service, item_id)
    else:
        files_to_process = [root_item]
    
    if not files_to_process:
        raise HTTPException(status_code=404, detail="No files found in the selection or its sub-folders.")

    all_documents = []
    
    # 3. Process the discovered files
    for f in files_to_process:
        file_id = f['id']
        file_name = f['name']
        mime_type = f['mimeType']
        text_content = ""
        
        try:
            if mime_type == 'application/vnd.google-apps.document':
                request = service.files().export_media(fileId=file_id, mimeType='text/plain')
                text_content = request.execute().decode('utf-8')
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                request = service.files().get_media(fileId=file_id)
                content = request.execute()
                import docx, io
                doc = docx.Document(io.BytesIO(content))
                text_content = "\n".join([para.text for para in doc.paragraphs])
            elif mime_type == 'application/pdf':
                request = service.files().get_media(fileId=file_id)
                content = request.execute()
                from PyPDF2 import PdfReader
                import io
                pdf_reader = PdfReader(io.BytesIO(content))
                for page in pdf_reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text_content += extracted + "\n"
            else:
                # Basic text/code file fallback
                request = service.files().get_media(fileId=file_id)
                content = request.execute()
                text_content = content.decode('utf-8', errors='ignore')
                
            if text_content and text_content.strip():
                all_documents.append({
                    "text": text_content, 
                    "metadata": {"source": file_name, "user_id": user_id}
                })
        except Exception as e:
            print(f"Skipping {file_name}: Unreadable format or {str(e)}")

    if not all_documents:
        raise HTTPException(status_code=400, detail="Could not extract readable text from the selected item(s).")

    # 4. Chunk and Save to Vector DB
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
    chunks = []
    metadatas = []
    
    for doc in all_documents:
        split_texts = text_splitter.split_text(doc["text"])
        chunks.extend(split_texts)
        metadatas.extend([doc["metadata"]] * len(split_texts))
        
    from langchain_chroma import Chroma
    if chunks:
        vector_store = Chroma.from_texts(
            texts=chunks, embedding=embeddings, metadatas=metadatas, persist_directory=CHROMA_PERSIST_DIR
        )
        return {
            "message": "Ingestion Complete!", 
            "files_processed": len(all_documents),
            "total_chunks_saved": len(chunks)
        }
    else:
        return {"message": "No text chunks were generated."}


# 6. GET ALL DRIVE ITEMS (For the React UI Modal)
import json
from fastapi.responses import StreamingResponse
import asyncio

# 6. STREAMING DRIVE ITEMS (Super Fast UI Updates)
@app.get("/list-drive-items/{user_id}")
async def list_drive_items(user_id: str):
    """Streams files and folders from Google Drive, newest first."""
    if user_id not in user_sessions:
        raise HTTPException(status_code=401, detail="User not authenticated")
        
    from google.oauth2.credentials import Credentials
    creds = Credentials(**user_sessions[user_id])
    service = build('drive', 'v3', credentials=creds)
    
    async def event_generator():
        page_token = None
        # Strict query to keep things light
        query = ("(mimeType='application/vnd.google-apps.folder' or "
                 "mimeType='application/vnd.google-apps.document' or "
                 "mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document' or "
                 "mimeType='application/pdf') and trashed=false")
        
        try:
            # We will fetch up to 10 pages, but stream each one immediately
            for _ in range(10): 
                # Run the blocking Google API call in a separate thread to keep streaming smooth
                loop = asyncio.get_event_loop()
                results = await loop.run_in_executor(None, lambda: service.files().list(
                    q=query,
                    corpora="allDrives",
                    includeItemsFromAllDrives=True,
                    supportsAllDrives=True,
                    fields="nextPageToken, files(id, name, mimeType, shared, ownedByMe, modifiedTime)",
                    orderBy="modifiedTime desc", # Fetch newest files first for better UX
                    pageSize=200, # Smaller pages = faster "First Byte" to UI
                    pageToken=page_token
                ).execute())
                
                items = results.get('files', [])
                page_token = results.get('nextPageToken')
                
                if items:
                    # Yield this chunk of files as a JSON string followed by a newline
                    yield json.dumps(items) + "\n"
                
                if not page_token:
                    break
        except Exception as e:
            yield json.dumps({"error": str(e)}) + "\n"

    return StreamingResponse(event_generator(), media_type="application/x-ndjson")